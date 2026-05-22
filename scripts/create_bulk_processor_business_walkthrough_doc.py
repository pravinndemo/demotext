from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT_PATH = Path(r"c:\dev\demotext\documents\bulk-processor-business-walkthrough-script.docx")

ACCENT = RGBColor(31, 78, 95)
MUTED = RGBColor(92, 107, 117)
BODY = RGBColor(34, 34, 34)


SECTIONS = [
    (
        "Opening",
        [
            "Today I’m going to give a short walk-through of the bulk processor and why it matters to the business.",
            "The simplest way to think about it is this: we have a repeatable way to take a large set of records, stage them safely, validate them, and then hand them off for processing without relying on manual record-by-record work. The design gives us control, visibility, and scale.",
        ],
    ),
    (
        "Why we built it",
        [
            "We needed a process that can handle large volumes consistently. The business problem is not just speed. It is also accuracy, traceability, and the ability to stop and review before anything final happens.",
            "If we processed every item manually, we would be slow and create more risk. If we pushed everything straight through, we would lose control and make recovery harder. This solution gives us a middle ground: users can prepare data in bulk, review it, and only then submit it for processing.",
        ],
    ),
    (
        "End-to-end flow",
        [
            "There are three parts to the journey. First, the user opens the Bulk Processor record in the model-driven app and either selects items or uploads a CSV file. Second, the user saves the items. That stages the records and keeps the batch in Draft. Third, when the batch is ready, the user submits it. That moves the batch to Queued and signals that background processing can start.",
            "A timer-driven worker then picks up queued batches and completes the downstream work. There is also a separate SVT tracking route for single-item dispatches, but the bulk journey is the main path we are talking about today.",
        ],
    ),
    (
        "What happens at Save Items",
        [
            "Save Items is the safety step. The purpose is to create or update staged child rows, not to complete the business process. The batch stays in Draft. The system validates the incoming data, updates counts on the parent record, and lets the user see how many rows are valid, invalid, duplicate, or still pending.",
            "If the data came from a file, the function reads the CSV from Dataverse. If it came from selection, it uses the selected SSU identifiers. One important rule is that missing rows are not treated as deletes. If a row is removed, that has to happen explicitly in the form or subgrid. That protects us from accidentally deleting rows just because they were not included in a payload.",
        ],
    ),
    (
        "What happens at Submit Batch",
        [
            "Submit Batch is the business decision point. By this stage the user has reviewed the staged rows and is ready to proceed. The function checks that the batch is still in Draft, that there is at least one item to process, and that there is at least one valid item.",
            "If those checks pass, the batch is moved to Queued. From a user perspective, that means the data has passed staging and is now ready for controlled background processing. In the current implementation, the actual request and job creation work is not done in the user’s click path; it is deferred to the worker so the submit action stays fast and reliable.",
        ],
    ),
    (
        "Background processing",
        [
            "The timer runs on a schedule and picks up queued batches. It works in chunks rather than trying to process everything at once. That matters because it reduces pressure on Dataverse, improves reliability, and gives us a clean place to retry transient failures.",
            "The timer updates the parent batch as it goes, and because the work is chunked, one failed piece does not necessarily stop the entire batch. That gives us partial success rather than all-or-nothing behavior. The system also recalculates counts from the records actually stored in Dataverse, so the final state is based on what really happened rather than what we hoped would happen.",
        ],
    ),
    (
        "Controls and resilience",
        [
            "There are a few important control points built into the design. The first is the status gate: only Draft batches can be staged or submitted. The second is validation: the system checks payload shape, required IDs, and batch eligibility. The third is observability: each request carries a correlation ID, and the processor stamps processing start, completion, and error details onto the batch.",
            "The fourth is retry handling: transient errors are retried, and chunk failures are isolated. These controls mean we can support the process safely in a live environment and troubleshoot issues without guessing.",
        ],
    ),
    (
        "Business outcome",
        [
            "The result is a controlled, repeatable bulk process. The business gets faster handling of large volumes, fewer manual steps, better data quality, and a much clearer audit trail.",
            "The system can absorb growth because work is staged and processed in controlled batches. And when something goes wrong, we can see exactly where it failed and what still succeeded. That makes the process more predictable for the business and easier for support to manage.",
        ],
    ),
    (
        "Close",
        [
            "So the headline is simple: users prepare the data, validate it, and submit it when ready. The platform handles the heavy lifting in the background, with controls that keep the process safe, scalable, and transparent.",
            "That is the main value of the bulk processor. If helpful, I can also walk through the entity model, the request routes, or the processing status flow in more detail.",
        ],
    ),
    (
        "Short summary",
        [
            "If I had to summarize it in one sentence, it is this: the bulk processor lets us stage high-volume work safely, validate it before it is final, and then hand it to the background worker in a way that is scalable and auditable.",
            "So the business keeps control, the platform handles the volume, and support gets a clear record of what happened at every stage.",
        ],
    ),
]


def set_font(run, name="Arial", size=12, bold=False, italic=False, color=BODY):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def add_bottom_border(paragraph, color="D9DEE3", size="6"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)


def add_page_number_field(paragraph):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_text = OxmlElement("w:t")
    fld_text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(fld_text)
    run._r.append(fld_end)


def format_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(12)
    normal.font.color.rgb = BODY
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08

    title_style = styles["Title"]
    title_style.font.name = "Arial"
    title_style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    title_style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    title_style.font.size = Pt(22)
    title_style.font.bold = True
    title_style.font.color.rgb = ACCENT

    for style_name, size in [("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 12)]:
        style = styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = ACCENT
        style.paragraph_format.space_before = Pt(12 if style_name == "Heading 1" else 8)
        style.paragraph_format.space_after = Pt(6 if style_name != "Heading 3" else 3)

    header = section.header
    header_p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header_p.paragraph_format.space_after = Pt(0)
    header_p.paragraph_format.space_before = Pt(0)
    header_run = header_p.add_run("Bulk Processor business walkthrough script")
    set_font(header_run, size=9, color=MUTED)
    add_bottom_border(header_p, color="D9DEE3", size="6")

    footer = section.footer
    footer_p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    footer_p.paragraph_format.space_before = Pt(0)
    footer_p.paragraph_format.space_after = Pt(0)
    footer_p.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
    left = footer_p.add_run("Bulk Processor business walkthrough script")
    set_font(left, size=9, color=MUTED)
    footer_p.add_run("\t")
    add_page_number_field(footer_p)


def add_paragraph(doc: Document, text: str, *, bold_lead: str | None = None) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.08
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if bold_lead:
        lead = p.add_run(bold_lead)
        set_font(lead, bold=True)
        body = p.add_run(text)
        set_font(body)
    else:
        run = p.add_run(text)
        set_font(run)


def main() -> None:
    doc = Document()
    format_document(doc)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("Bulk Processor Business Walkthrough Script")
    set_font(title_run, size=22, bold=True, color=ACCENT)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(12)
    subtitle.paragraph_format.line_spacing = 1.08
    subtitle_run = subtitle.add_run("5 to 10 minute read-aloud version for a business call")
    set_font(subtitle_run, size=12, italic=True, color=MUTED)

    intro = doc.add_paragraph()
    intro.alignment = WD_ALIGN_PARAGRAPH.LEFT
    intro.paragraph_format.space_after = Pt(12)
    intro.paragraph_format.line_spacing = 1.08
    intro_text = (
        "How to use this script: read it at a steady pace, pause briefly between sections, "
        "and keep the focus on the business problem, the staged flow, and the control points. "
        "If you are short on time, you can trim the Background processing section and keep the Business outcome section."
    )
    intro_run = intro.add_run(intro_text)
    set_font(intro_run, italic=True, color=MUTED)

    for heading, paragraphs in SECTIONS:
        h = doc.add_paragraph(style="Heading 1")
        h.alignment = WD_ALIGN_PARAGRAPH.LEFT
        h_run = h.add_run(heading)
        set_font(h_run, size=16, bold=True, color=ACCENT)
        for idx, paragraph in enumerate(paragraphs):
            if idx == 0 and heading in {"Opening", "Why we built it"}:
                add_paragraph(doc, paragraph)
            elif idx == 0:
                add_paragraph(doc, paragraph)
            else:
                add_paragraph(doc, paragraph)

    doc.save(OUTPUT_PATH)


if __name__ == "__main__":
    main()
